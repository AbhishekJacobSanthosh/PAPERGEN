"""
AI Research Paper Generator v3.0 - Main Application
Production-ready Flask application with:
- Input validation and sanitization
- Error handling decorator
- Proper logging
- Security improvements
- User data support
"""
from flask import Flask, request, jsonify, render_template, send_file, Response, stream_with_context
from datetime import datetime
from functools import wraps
import os
import logging
import json
import glob
from functools import wraps
from flask import jsonify

# Import configuration
from config.settings import (
    SECRET_KEY, DEBUG, HOST, PORT, MAX_CONTENT_LENGTH,
    UPLOAD_FOLDER, SAVED_PAPERS_DIR
)

# Import models
from models.paper_structure import ResearchPaper, Author, Reference, Figure

# Import services
from services.paper_generator import PaperGeneratorService
from services.presentation_generator import PresentationGeneratorService
from services.rag_service import RAGService
from services.export_service import ExportService
from services.ocr_service import OCRService
from services.integrity_service import ContentIntegrityService

# Import utilities
from utils.text_processing import TextProcessor

# ==================== LOGGING SETUP ====================

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('paper_generator.log'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

# ==================== VALIDATION ====================

class ValidationError(Exception):
    """Custom validation error"""
    pass

class RequestValidator:
    """Validate incoming requests"""
    
    @staticmethod
    def validate_paper_generation(data: dict) -> None:
        """Validate paper generation request"""
        if not data.get('topic'):
            raise ValidationError("Topic is required")
        
        topic = data['topic'].strip()
        if len(topic) < 5:
            raise ValidationError("Topic must be at least 5 characters")
        if len(topic) > 2000:
            raise ValidationError("Topic must not exceed 500 characters")
        
        # Validate authors
        authors = data.get('authors', [])
        if not authors or len(authors) == 0:
            raise ValidationError("At least one author is required")
        
        for idx, author in enumerate(authors):
            if not author.get('name') or not author.get('email') or not author.get('affiliation'):
                raise ValidationError(f"Author {idx + 1} is missing required fields")
            
            # Basic email validation
            email = author['email']
            if '@' not in email or '.' not in email:
                raise ValidationError(f"Invalid email for author {idx + 1}")
        
        # Validate user data if provided
        user_data = data.get('user_data')
        if user_data:
            total_chars = sum(len(str(v)) for v in user_data.values() if v)
            if total_chars > 10000:
                raise ValidationError("User data exceeds maximum size (10,000 characters)")
    
    @staticmethod
    def validate_survey_request(data: dict) -> None:
        """Validate literature survey request"""
        papers = data.get('papers', [])
        if not papers:
            raise ValidationError("No papers provided")
        
        if len(papers) > 20:
            raise ValidationError("Too many papers (max 20)")
        
        topic = data.get('topic', '').strip()
        if len(topic) < 3:
            raise ValidationError("Topic must be at least 3 characters")
    
    @staticmethod
    def validate_retrieve_papers(data: dict) -> None:
        """Validate paper retrieval request"""
        topic = data.get('topic', '').strip()
        if not topic:
            raise ValidationError("Topic is required")
        
        if len(topic) < 3:
            raise ValidationError("Topic must be at least 3 characters")
        
        count = data.get('count', 7)
        if not isinstance(count, int) or count < 1 or count > 20:
            raise ValidationError("Count must be between 1 and 20")

# ==================== ERROR HANDLER DECORATOR ====================

def handle_api_errors(f):
    """Decorator to handle API errors consistently"""
    @wraps(f)
    def decorated_function(*args, **kwargs):
        try:
            return f(*args, **kwargs)
        except ValidationError as e:
            logger.warning(f"Validation error in {f.__name__}: {e}")
            return jsonify({'success': False, 'error': str(e)}), 400
        except Exception as e:
            logger.error(f"Error in {f.__name__}: {e}", exc_info=True)
            return jsonify({'success': False, 'error': str(e)}), 500
    return decorated_function

# ==================== FLASK APP SETUP ====================

app = Flask(__name__)
app.config['SECRET_KEY'] = SECRET_KEY
app.config['MAX_CONTENT_LENGTH'] = MAX_CONTENT_LENGTH
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

# Initialize services
paper_generator = PaperGeneratorService()
presentation_generator = PresentationGeneratorService()
rag_service = RAGService()
export_service = ExportService()
ocr_service = OCRService()
text_processor = TextProcessor()
integrity_service = ContentIntegrityService()

logger.info("="*70)
logger.info("🎓 AI Research Paper Generator v3.0")
logger.info("="*70)
logger.info("✓ All services initialized")
logger.info("="*70)

# ==================== ROUTES ====================

@app.route('/')
def home():
    """Serve the main page"""
    return render_template('index.html')

# ==================== PAPER GENERATION ====================

@app.route('/api/warmup', methods=['POST'])
@handle_api_errors
def warmup():
    """Warmup the LLM model"""
    success = paper_generator.llm.warmup()
    if success:
        logger.info("LLM warmup successful")
        return jsonify({'success': True, 'message': 'Model ready'})
    else:
        logger.warning("LLM warmup failed")
        return jsonify({'success': False, 'error': 'Warmup failed'}), 500
    
@app.route('/api/generate-titles', methods=['POST'])
@handle_api_errors
def generate_titles():
    """Generate multiple title options from description"""
    data = request.get_json()
    
    # Validate input
    description = data.get('description', '').strip()
    if not description:
        return jsonify({
            'success': False,
            'error': 'Description is required'
        }), 400
    
    if len(description) < 10:
        return jsonify({
            'success': False,
            'error': 'Description too short (minimum 10 characters)'
        }), 400
    
    if len(description) > 1000:
        return jsonify({
            'success': False,
            'error': 'Description too long (maximum 1000 characters)'
        }), 400
    
    count = data.get('count', 3)
    if not isinstance(count, int) or count < 2 or count > 5:
        count = 3
    
    logging.info(f"Generating {count} title options for description: {description[:100]}...")
    
    try:
        # Generate title options
        from models.llm import LLMInterface
        llm = LLMInterface()
        titles = llm.generate_title_options(description, count=count)
        
        if not titles or len(titles) == 0:
            return jsonify({
                'success': False,
                'error': 'Failed to generate titles'
            }), 500
        
        logging.info(f"Generated {len(titles)} title options successfully")
        
        return jsonify({
            'success': True,
            'titles': titles,
            'count': len(titles)
        })
        
    except Exception as e:
        logging.error(f"Error generating titles: {str(e)}")
        return jsonify({
            'success': False,
            'error': f'Title generation failed: {str(e)}'
        }), 500


@app.route('/api/generate-paper', methods=['POST'])
@handle_api_errors
def generate_paper_endpoint():
    """Generate complete research paper with optional user data"""
    data = request.json
    
    # Validate request
    RequestValidator.validate_paper_generation(data)
    
    topic = data.get('topic', '').strip()
    authors_data = data.get('authors', [])
    use_rag = data.get('use_rag', True)
    user_data = data.get('user_data')

    # Get selected title if any
    selected_title = data.get('selected_title', '').strip()
    
    logger.info(f"Generating paper - Topic: {topic[:50]}..., Authors: {len(authors_data)}, RAG: {use_rag}, Selected Title: {selected_title[:50]}")

    # Parse authors
    authors = []
    for author_data in authors_data:
        authors.append(Author(
            name=author_data.get('name', 'Unknown'),
            email=author_data.get('email', 'email@edu'),
            affiliation=author_data.get('affiliation', 'University')
        ))
    
    # Determine actual title to use
    if selected_title:
        paper_title = selected_title
        logger.info(f"Using user-selected title: {paper_title}")
    else:
        # Fallback: you may have a method to generate title from topic
        paper_title = paper_generator.generate_title(topic)  # Assuming this method exists
        logger.info(f"Generated title internally: {paper_title}")
    
    # Generate paper with user data and explicit title
    paper = paper_generator.generate_paper(topic, authors, use_rag, user_data, title=paper_title)
    
    logger.info(f"Paper generated successfully - Words: {paper.metadata['total_words']}")
    
    # Auto-save the paper
    try:
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"paper_{timestamp}.json"
        filepath = os.path.join(SAVED_PAPERS_DIR, filename)
        
        with open(filepath, 'w') as f:
            json.dump(paper.to_dict(), f, indent=4)
            
        logger.info(f"Auto-saved paper to {filepath}")
    except Exception as e:
        logger.error(f"Failed to auto-save paper: {e}")

    return jsonify({
        'success': True,
        'paper': paper.to_dict()
    })


@app.route('/api/generate-paper-stream', methods=['POST'])
def generate_paper_stream_endpoint():
    """Generate research paper with streaming updates (SSE)"""
    data = request.json
    
    # Validate request (reuse existing validator logic manually or via try-catch)
    try:
        RequestValidator.validate_paper_generation(data)
    except ValidationError as e:
        return jsonify({'success': False, 'error': str(e)}), 400
        
    topic = data.get('topic', '').strip()
    authors_data = data.get('authors', [])
    use_rag = data.get('use_rag', True)
    user_data = data.get('user_data')
    selected_title = data.get('selected_title', '').strip()
    
    logger.info(f"Streaming paper generation - Topic: {topic[:50]}...")
    
    # Parse authors
    authors = []
    for author_data in authors_data:
        authors.append(Author(
            name=author_data.get('name', 'Unknown'),
            email=author_data.get('email', 'email@edu'),
            affiliation=author_data.get('affiliation', 'University')
        ))
        
    def generate():
        # Generator wrapper for SSE
        try:
            # Determine title
            paper_title = selected_title if selected_title else None
            
            # Call streaming service
            for event_json in paper_generator.generate_paper_stream(
                topic, authors, use_rag, user_data, title=paper_title
            ):
                yield f"data: {event_json}\n\n"
                
        except Exception as e:
            logger.error(f"Streaming error: {e}", exc_info=True)
            error_json = json.dumps({'status': 'error', 'message': str(e)})
            yield f"data: {error_json}\n\n"

    return Response(stream_with_context(generate()), mimetype='text/event-stream')


@app.route('/api/latest-paper', methods=['GET'])
def get_latest_paper():
    """Retrieve the most recently saved paper"""
    try:
        list_of_files = glob.glob(os.path.join(SAVED_PAPERS_DIR, '*.json'))
        if not list_of_files:
            return jsonify({'success': False, 'error': 'No saved papers found'}), 404
            
        latest_file = max(list_of_files, key=os.path.getctime)
        
        with open(latest_file, 'r') as f:
            paper_data = json.load(f)
            
        logger.info(f"Retrieved latest paper: {latest_file}")
        return jsonify({'success': True, 'paper': paper_data})
    except Exception as e:
        logger.error(f"Error retrieving latest paper: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500


# ==================== LITERATURE SURVEY ====================

@app.route('/api/retrieve-papers', methods=['POST'])
@handle_api_errors
def retrieve_papers():
    """Retrieve papers from Semantic Scholar"""
    data = request.json
    
    # Validate request
    RequestValidator.validate_retrieve_papers(data)
    
    topic = data.get('topic', '').strip()
    count = data.get('count', 7)
    
    logger.info(f"Retrieving papers - Topic: {topic}, Count: {count}")
    
    papers = rag_service.search_papers(topic, limit=count)
    
    papers_json = [
        {
            'title': p.title,
            'authors': p.authors,
            'year': p.year,
            'venue': p.venue,
            'citationCount': p.citation_count,
            'doi': p.doi,
            'url': p.url,
            'abstract': p.abstract
        } for p in papers
    ]
    
    logger.info(f"Retrieved {len(papers_json)} papers")
    
    return jsonify({
        'success': True,
        'papers': papers_json,
        'count': len(papers_json)
    })

@app.route('/api/generate-survey', methods=['POST'])
@handle_api_errors
def generate_survey():
    """Generate literature survey from papers"""
    data = request.json
    
    # Validate request
    RequestValidator.validate_survey_request(data)
    
    papers_data = data.get('papers', [])
    topic = data.get('topic', 'the research area')
    
    logger.info(f"Generating survey - Topic: {topic}, Papers: {len(papers_data)}")
    
    # Build context
    context = "Research papers retrieved:\n\n"
    for i, paper in enumerate(papers_data, 1):
        authors_str = ', '.join(paper.get('authors', ['Unknown'])[:3])
        if len(paper.get('authors', [])) > 3:
            authors_str += ' et al.'
        
        context += f"Paper {i}:\n"
        context += f"Title: {paper.get('title')}\n"
        context += f"Authors: {authors_str}\n"
        context += f"Year: {paper.get('year')}\n"
        context += f"Citations: {paper.get('citationCount', 0)}\n"
        context += f"Abstract: {paper.get('abstract', 'No abstract')[:250]}\n\n"
    
    prompt = f"""Write a comprehensive literature survey on "{topic}" based on the research papers provided above.

Structure your survey with clear sections as follows:

Introduction
Write 2-3 paragraphs introducing the research area "{topic}", its importance, and the scope of this survey.

Summary of Key Papers and Their Contributions
For each major paper on "{topic}", describe the research objectives, methodology, key findings, and contributions to the field. Write in continuous prose paragraphs.

Common Themes and Approaches
Identify and discuss the common methodologies, techniques, and approaches used across the papers in "{topic}" research. Group related work together.

Research Gaps and Opportunities
Discuss what has not been addressed in "{topic}" research, limitations of existing work, and potential future research directions.

Conclusion
Summarize the state of the "{topic}" field and key takeaways from the literature.

ABSOLUTE REQUIREMENTS:
- Write ONLY in plain text prose paragraphs
- NO markdown - no #, *, **, _, or bullet points
- Start each section name on its own line
- Separate paragraphs with double line breaks
- Use formal academic language
- Target 800 - 1000 words total
- Reference specific papers by author names and years
- Write continuously - no lists or numbered items
- Always use the complete topic name "{topic}" when referring to the field

Begin the literature survey:"""

    survey = paper_generator.llm.generate(
        prompt,
        temperature=0.7,
        max_tokens=1400,
        context=context
    )
    
    if survey:
        # Clean the generated survey
        survey = text_processor.clean_survey_text(survey)
        logger.info(f"Survey generated - Words: {len(survey.split())}")
        return jsonify({'success': True, 'survey': survey, 'title': f'Literature Survey: {topic}'})
    else:
        logger.error("Survey generation failed")
        return jsonify({'success': False, 'error': 'Generation failed'}), 500

@app.route('/api/download-survey-pdf', methods=['POST'])
@handle_api_errors
def download_survey_pdf():
    """Export literature survey as professional PDF"""
    data = request.json
    survey_text = data.get('survey', '')
    topic = data.get('topic', 'Literature Survey')
    
    if not survey_text:
        raise ValidationError("No survey content")
    
    logger.info(f"Generating survey PDF - Topic: {topic}")
    
    from reportlab.lib.pagesizes import letter # pyright: ignore[reportMissingModuleSource]
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle # pyright: ignore[reportMissingModuleSource]
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer # pyright: ignore[reportMissingModuleSource]
    from reportlab.lib.units import inch # pyright: ignore[reportMissingModuleSource]
    from reportlab.lib.enums import TA_JUSTIFY, TA_CENTER, TA_LEFT # pyright: ignore[reportMissingModuleSource]
    from reportlab.lib import colors # pyright: ignore[reportMissingModuleSource]
    from reportlab.platypus import HRFlowable # pyright: ignore[reportMissingModuleSource]
    from io import BytesIO
    
    # Clean the survey text thoroughly
    survey_text = text_processor.clean_survey_text(survey_text)
    
    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer, 
        pagesize=letter,
        leftMargin=1.25*inch,
        rightMargin=1.25*inch,
        topMargin=1*inch,
        bottomMargin=1*inch
    )
    
    styles = getSampleStyleSheet()
    
    # Title style
    title_style = ParagraphStyle(
        'SurveyTitle',
        parent=styles['Title'],
        fontSize=22,
        alignment=TA_CENTER,
        spaceAfter=30,
        spaceBefore=10,
        fontName='Times-Bold',
        textColor=colors.HexColor('#1a1a1a')
    )
    
    # Subtitle style
    subtitle_style = ParagraphStyle(
        'Subtitle',
        parent=styles['Normal'],
        fontSize=12,
        alignment=TA_CENTER,
        spaceAfter=40,
        fontName='Times-Italic',
        textColor=colors.HexColor('#666666')
    )
    
    # Section heading style
    heading_style = ParagraphStyle(
        'SectionHeading',
        parent=styles['Heading2'],
        fontSize=14,
        alignment=TA_LEFT,
        spaceAfter=12,
        spaceBefore=20,
        fontName='Times-Bold',
        textColor=colors.HexColor('#2c3e50'),
        leftIndent=0
    )
    
    # Body text style
    body_style = ParagraphStyle(
        'SurveyBody',
        parent=styles['BodyText'],
        fontSize=11,
        leading=18,
        alignment=TA_JUSTIFY,
        spaceAfter=12,
        fontName='Times-Roman',
        textColor=colors.HexColor('#2c2c2c')
    )
    
    story = []
    
    # Title page
    story.append(Spacer(1, 1*inch))
    story.append(Paragraph(topic, title_style))
    story.append(Paragraph(f"Generated on {datetime.now().strftime('%B %d, %Y')}", subtitle_style))
    story.append(Spacer(1, 0.5*inch))
    
    # Horizontal line
    story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor('#cccccc')))
    story.append(Spacer(1, 0.3*inch))
    
    # Process the survey text
    paragraphs = survey_text.split('\n\n')
    
    section_keywords = [
        'Introduction', 'Background', 'Overview',
        'Summary', 'Literature Review', 'Key Papers', 'Contributions',
        'Common Themes', 'Approaches', 'Methodologies', 'Methods',
        'Research Gaps', 'Challenges', 'Opportunities', 'Limitations',
        'Discussion', 'Conclusion', 'Future Work', 'Future Directions'
    ]
    
    for para in paragraphs:
        if not para.strip():
            continue
        
        # Check if this paragraph is a section header
        is_heading = False
        para_clean = para.strip()
        
        # Check if starts with section keyword and is short
        for keyword in section_keywords:
            if para_clean.startswith(keyword) and len(para_clean) < 100:
                is_heading = True
                break
        
        if is_heading:
            story.append(Paragraph(para_clean, heading_style))
        else:
            # Regular body paragraph
            story.append(Paragraph(para_clean, body_style))
    
    # Build the PDF
    doc.build(story)
    buffer.seek(0)
    
    logger.info("Survey PDF generated successfully")
    
    return send_file(
        buffer,
        as_attachment=True,
        download_name=f"literature_survey_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf",
        mimetype='application/pdf'
    )

@app.route('/api/download-survey-docx', methods=['POST'])
@handle_api_errors
def download_survey_docx():
    """Export literature survey as DOCX"""
    data = request.json
    survey_text = data.get('survey', '')
    topic = data.get('topic', 'Literature Survey')
    
    if not survey_text:
        raise ValidationError("No survey content")
    
    logger.info(f"Generating survey DOCX - Topic: {topic}")
    
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from io import BytesIO
    
    # Clean the survey text
    survey_text = text_processor.clean_survey_text(survey_text)
    
    doc = Document()
    
    # Title
    title = doc.add_heading(topic, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Date
    date_para = doc.add_paragraph(f"Generated on {datetime.now().strftime('%B %d, %Y')}")
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.runs[0]
    date_run.font.size = Pt(10)
    date_run.font.italic = True
    date_run.font.color.rgb = RGBColor(100, 100, 100)
    
    doc.add_paragraph()
    
    # Survey content
    paragraphs = survey_text.split('\n\n')
    
    section_keywords = [
        'Introduction', 'Background', 'Overview',
        'Summary', 'Literature Review', 'Key Papers', 'Contributions',
        'Common Themes', 'Approaches', 'Methodologies', 'Methods',
        'Research Gaps', 'Challenges', 'Opportunities', 'Limitations',
        'Discussion', 'Conclusion', 'Future Work', 'Future Directions'
    ]
    
    for para_text in paragraphs:
        if not para_text.strip():
            continue
        
        para_clean = para_text.strip()
        
        # Check if this is a section header
        is_heading = False
        for keyword in section_keywords:
            if para_clean.startswith(keyword) and len(para_clean) < 100:
                is_heading = True
                break
        
        if is_heading:
            doc.add_heading(para_clean, level=1)
        else:
            para = doc.add_paragraph(para_clean)
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    logger.info("Survey DOCX generated successfully")
    
    return send_file(
        buffer,
        as_attachment=True,
        download_name=f"literature_survey_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

# ==================== EXPORT ====================

@app.route('/api/download-pdf', methods=['POST'])
@handle_api_errors
def download_pdf():
    """Export paper as PDF"""
    data = request.json
    paper_data = data.get('paper', {})
    
    if not paper_data:
        raise ValidationError("No paper data")
    
    logger.info(f"Generating paper PDF - Title: {paper_data.get('title', 'Unknown')[:50]}...")
    
    # Reconstruct paper object
    from models.paper_structure import Figure
    
    paper = ResearchPaper(
        title=paper_data['title'],
        authors=[Author(**a) for a in paper_data['authors']],
        abstract=paper_data['abstract'],
        sections=paper_data['sections'],
        references=[Reference(**r) for r in paper_data.get('references', [])],
        figures={k: Figure(**v) for k, v in paper_data.get('figures', {}).items()},
        doi=paper_data['doi'],
        generated_at=datetime.fromisoformat(paper_data['generated_at'])
    )
    
    buffer = export_service.generate_pdf(paper)
    
    logger.info("Paper PDF generated successfully")
    
    return send_file(
        buffer,
        as_attachment=True,
        download_name=f"paper_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf",
        mimetype='application/pdf'
    )

@app.route('/api/download-docx', methods=['POST'])
@handle_api_errors
def download_docx():
    """Export paper as DOCX"""
    data = request.json
    paper_data = data.get('paper', {})
    
    if not paper_data:
        raise ValidationError("No paper data")
    
    logger.info(f"Generating paper DOCX - Title: {paper_data.get('title', 'Unknown')[:50]}...")
    
    # Reconstruct paper object
    paper = ResearchPaper(
        title=paper_data['title'],
        authors=[Author(**a) for a in paper_data['authors']],
        abstract=paper_data['abstract'],
        sections=paper_data['sections'],
        references=[Reference(**r) for r in paper_data.get('references', [])],
        figures={},
        doi=paper_data['doi'],
        generated_at=datetime.fromisoformat(paper_data['generated_at'])
    )
    
    buffer = export_service.generate_docx(paper)
    
    logger.info("Paper DOCX generated successfully")
    
    return send_file(
        buffer,
        as_attachment=True,
        download_name=f"paper_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

# ==================== OCR ====================

@app.route('/api/extract-ocr', methods=['POST'])
@handle_api_errors
def extract_ocr():
    """Extract text from uploaded image"""
    if 'image' not in request.files:
        raise ValidationError("No image provided")
    
    file = request.files['image']
    if not file.filename:
        raise ValidationError("Empty filename")
    
    logger.info(f"Processing OCR - File: {file.filename}")
    
    # Save temporarily
    import uuid
    filename = f"ocr_{uuid.uuid4().hex[:8]}.jpg"
    filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    file.save(filepath)
    
    # Extract text
    text = ocr_service.extract_text(filepath)
    
    # Clean up
    try:
        os.remove(filepath)
    except:
        pass
    
    logger.info(f"OCR extracted {len(text)} characters")
    
    return jsonify({'success': True, 'text': text})

# ==================== ERROR HANDLERS ====================

@app.errorhandler(404)
def not_found(error):
    """Handle 404 errors"""
    logger.warning(f"404 error: {request.url}")
    return jsonify({'success': False, 'error': 'Endpoint not found'}), 404

@app.errorhandler(500)
def internal_error(error):
    """Handle 500 errors"""
    logger.error(f"500 error: {error}")
    return jsonify({'success': False, 'error': 'Internal server error'}), 500

# ==================== MAIN ====================

@app.route('/api/download-pptx', methods=['POST'])
def download_pptx():
    """Generate and download PPTX"""
    data = request.json
    if not data or 'paper' not in data:
        return jsonify({'success': False, 'error': 'No paper data provided'}), 400
        
    try:
        # Reconstruct paper object
        paper_data = data['paper']
        paper = ResearchPaper(
            title=paper_data.get('title', 'Untitled'),
            authors=[Author(**a) for a in paper_data.get('authors', [])],
            abstract=paper_data.get('abstract', ''),
            sections=paper_data.get('sections', {}),
            references=[Reference(**r) for r in paper_data.get('references', [])],
            figures={}, # Figures not supported in PPTX yet
            doi=paper_data.get('doi'),
            generated_at=datetime.now()
        )
        
        # Generate PPTX
        filename = f"presentation_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pptx"
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        presentation_generator.generate_presentation(paper, filepath)
        
        return send_file(
            filepath,
            as_attachment=True,
            download_name=f"{paper.title[:30]}_presentation.pptx",
            mimetype='application/vnd.openxmlformats-officedocument.presentationml.presentation'
        )
        
    except Exception as e:
        logger.error(f"PPTX generation error: {e}", exc_info=True)
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/check-integrity', methods=['POST'])
def check_integrity():
    """Check paper integrity (plagiarism and AI detection)"""
    try:
        data = request.json
        paper_data = data.get('paper')
        
        if not paper_data:
            return jsonify({'error': 'No paper data provided'}), 400
            
        # Reconstruct paper content
        full_text = f"{paper_data.get('title', '')}\n\n{paper_data.get('abstract', '')}\n\n"
        for section in paper_data.get('sections', {}).values():
            full_text += f"{section}\n\n"
            
        # Get source documents (abstracts from references)
        # Note: In a real scenario, we'd want the full text of references, 
        # but here we use what we have (abstracts if available in RAG cache)
        # For this implementation, we'll use the reference titles/abstracts passed in paper_data
        source_docs = []
        for ref in paper_data.get('references', []):
            if isinstance(ref, dict):
                source_docs.append(f"{ref.get('title', '')} {ref.get('abstract', '')}")
        
        # Run checks
        plagiarism_result = integrity_service.check_plagiarism(full_text, source_docs)
        ai_result = integrity_service.detect_ai_content(full_text)
        
        return jsonify({
            'plagiarism': plagiarism_result,
            'ai_detection': ai_result
        })
        
    except Exception as e:
        logger.error(f"Integrity check failed: {e}")
        return jsonify({'error': str(e)}), 500

if __name__ == '__main__':
    logger.info(f"🚀 Starting server on http://{HOST}:{PORT}")
    logger.info(f"📝 Model: {paper_generator.llm.model}")
    logger.info("="*70)
    
    app.run(debug=DEBUG, host=HOST, port=PORT)
