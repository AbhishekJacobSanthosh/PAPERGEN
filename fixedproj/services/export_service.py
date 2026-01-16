"""
Export Service - PDF and DOCX generation
"""
from io import BytesIO
from datetime import datetime
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.enums import TA_JUSTIFY, TA_CENTER
from reportlab.lib import colors
from reportlab.platypus import (
    BaseDocTemplate, PageTemplate, Frame, Paragraph, 
    Spacer, Table, TableStyle, Image as RLImage,
    FrameBreak, NextPageTemplate
)
import base64
from models.paper_structure import ResearchPaper

class ExportService:
    """Service for exporting papers to PDF and DOCX"""
    
    def _fix_encoding(self, text: str) -> str:
        """Fix common UTF-8 encoding issues in text"""
        if not text:
            return text
        encoding_fixes = {
            'â€¢': '•', 'Â²': '²', 'Â³': '³',
            'Ã©': 'é', 'Ã¡': 'á', 'Ã­': 'í', 'Ã³': 'ó', 'Ãº': 'ú',
            'Ã±': 'ñ', 'Ã¼': 'ü', 'Ã¶': 'ö', 'Ã¤': 'ä',
            'â€"': '—', 'â€™': "'", 'â€œ': '"', 'â€': '"',
        }
        result = text
        for bad, good in encoding_fixes.items():
            result = result.replace(bad, good)
        return result
    
    def generate_pdf(self, paper: ResearchPaper) -> BytesIO:
        """Generate IEEE-format PDF with correct layout"""
        buffer = BytesIO()
        
        # Setup document
        doc = BaseDocTemplate(
            buffer,
            pagesize=letter,
            leftMargin=0.75*inch,
            rightMargin=0.75*inch,
            topMargin=0.75*inch,
            bottomMargin=0.75*inch
        )
        
        # Dimensions
        full_width = letter[0] - 1.5*inch
        col_width = (full_width - 0.2*inch) / 2
        full_height = letter[1] - 1.5*inch
        
        # --- First Page Layout ---
        # Top Frame for Title/Authors
        header_height = 2.5*inch 
        body_height = full_height - header_height - 0.2*inch
        
        frame_top = Frame(
            0.75*inch, 
            0.75*inch + body_height + 0.2*inch, 
            full_width, 
            header_height, 
            id='top',
            showBoundary=0
        )
        
        frame_left_1 = Frame(
            0.75*inch, 
            0.75*inch, 
            col_width, 
            body_height, 
            id='left_1',
            showBoundary=0
        )
        
        frame_right_1 = Frame(
            0.75*inch + col_width + 0.2*inch, 
            0.75*inch, 
            col_width, 
            body_height, 
            id='right_1',
            showBoundary=0
        )
        
        # --- Later Pages Layout ---
        frame_left_2 = Frame(
            0.75*inch, 
            0.75*inch, 
            col_width, 
            full_height, 
            id='left_2',
            showBoundary=0
        )
        
        frame_right_2 = Frame(
            0.75*inch + col_width + 0.2*inch, 
            0.75*inch, 
            col_width, 
            full_height, 
            id='right_2',
            showBoundary=0
        )
        
        # Templates
        template_first = PageTemplate(id='FirstPage', frames=[frame_top, frame_left_1, frame_right_1])
        template_later = PageTemplate(id='LaterPage', frames=[frame_left_2, frame_right_2])
        
        doc.addPageTemplates([template_first, template_later])
        
        # Define styles
        styles = self._create_styles()
        
        # Build content
        elements = []
        
        # --- Header Content (Flows into 'top' frame) ---
        
        # Title
        elements.append(Paragraph(paper.title, styles['pdf_title']))
        elements.append(Spacer(1, 0.1*inch))
        
        # Authors
        if paper.authors:
            author_data = []
            row = []
            for i, author in enumerate(paper.authors):
                block = [
                    Paragraph(f"<b>{author.name}</b>", styles['pdf_author_name']),
                    Paragraph(f"<i>{author.affiliation}</i>", styles['pdf_author_affil']),
                    Paragraph(f"{author.email}", styles['pdf_author_email'])
                ]
                row.append(block)
                if len(row) == 3 or i == len(paper.authors) - 1:
                    author_data.append(row)
                    row = []
            
            num_cols = len(author_data[0])
            # Use full width for author table
            t_col_width = full_width / num_cols
            
            t = Table(author_data, colWidths=[t_col_width] * num_cols)
            t.setStyle(TableStyle([
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ]))
            elements.append(t)
            elements.append(Spacer(1, 0.1*inch))
            
        # Force break to left column of first page
        elements.append(FrameBreak())
        
        # Switch to LaterPage template for subsequent pages
        elements.append(NextPageTemplate('LaterPage'))
        
        # --- Body Content (Flows into 'left_1', 'right_1', then 'left_2', 'right_2'...) ---
        
        # Abstract
        abstract_text = f"<b><i>Abstract</i></b>—{paper.abstract}"
        elements.append(Paragraph(abstract_text, styles['pdf_abstract']))
        elements.append(Spacer(1, 0.1*inch))
        
        # DOI
        if paper.doi:
            elements.append(Paragraph(f"<b>DOI:</b> {paper.doi}", styles['pdf_body']))
            elements.append(Spacer(1, 0.15*inch))
        
        # Sections
        section_titles = {
            'introduction': 'I. INTRODUCTION',
            'literature_review': 'II. LITERATURE REVIEW',
            'methodology': 'III. METHODOLOGY',
            'results': 'IV. RESULTS',
            'discussion': 'V. DISCUSSION',
            'conclusion': 'VI. CONCLUSION',
            'references': 'REFERENCES'
        }
        
        for key, title in section_titles.items():
            if key in paper.sections:
                elements.append(Paragraph(title, styles['pdf_section']))
                
                content = paper.sections[key]
                # Convert single newlines to <br/> for proper line breaks in PDF
                # First, normalize multiple newlines, then convert single newlines
                formatted_content = content.replace('\n\n', '<<PARA>>').replace('\n', '<br/>').replace('<<PARA>>', '\n\n')
                
                for para in formatted_content.split('\n\n'):
                    if para.strip():
                        # Clean up any remaining formatting issues
                        clean_para = para.strip()
                        elements.append(Paragraph(clean_para, styles['pdf_body']))
                
                if key == 'results':
                    elements.extend(self._add_figures_to_pdf(paper.figures, col_width, styles))
                
                elements.append(Spacer(1, 0.08*inch))
        
        # Build PDF
        doc.build(elements)
        buffer.seek(0)
        
        return buffer

    def generate_docx(self, paper: ResearchPaper) -> BytesIO:
        """Generate DOCX document"""
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            doc = Document()
            
            # Title
            title = doc.add_heading(paper.title, 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Authors
            for author in paper.authors:
                author_para = doc.add_paragraph()
                author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                name_run = author_para.add_run(f"{author.name}\n")
                name_run.font.size = Pt(10)
                
                affil_run = author_para.add_run(f"{author.affiliation}\n")
                affil_run.font.size = Pt(9)
                affil_run.font.italic = True
                
                email_run = author_para.add_run(author.email)
                email_run.font.size = Pt(9)
            
            doc.add_paragraph()
            
            # Abstract
            abstract_heading = doc.add_paragraph()
            heading_run = abstract_heading.add_run('Abstract')
            heading_run.font.bold = True
            heading_run.font.italic = True
            
            abstract_para = doc.add_paragraph(paper.abstract)
            abstract_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            # Sections
            section_titles = {
                'introduction': 'I. INTRODUCTION',
                'literature_review': 'II. LITERATURE REVIEW',
                'methodology': 'III. METHODOLOGY',
                'results': 'IV. RESULTS',
                'discussion': 'V. DISCUSSION',
                'conclusion': 'VI. CONCLUSION',
                'references': 'REFERENCES'
            }
            
            for key, title in section_titles.items():
                if key in paper.sections:
                    doc.add_heading(title, level=1)
                    
                    content = paper.sections[key]
                    for para_text in content.split('\n\n'):
                        if para_text.strip():
                            body_para = doc.add_paragraph(para_text.strip())
                            body_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            # Save
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            
            return buffer
            
        except ImportError:
            raise Exception("python-docx not installed. Run: pip install python-docx")

    def generate_html(self, paper: ResearchPaper) -> str:
        """Generate HTML with proper CSS formatting for structured content"""
        
        # CSS for IEEE-like formatting with two columns
        css = """
        <style>
            @media print { body { font-size: 10pt; } }
            body { font-family: 'Times New Roman', serif; max-width: 900px; margin: 0 auto; padding: 20px; line-height: 1.4; font-size: 10pt; }
            h1 { text-align: center; font-size: 22px; margin-bottom: 10px; column-span: all; }
            .header { column-span: all; text-align: center; margin-bottom: 15px; }
            .authors { text-align: center; margin-bottom: 15px; column-span: all; }
            .author { display: inline-block; margin: 0 10px; text-align: center; font-size: 9pt; }
            .author-name { font-weight: bold; }
            .author-affil { font-style: italic; }
            .author-email { font-size: 8pt; }
            .abstract { text-align: justify; margin-bottom: 15px; font-style: italic; column-span: all; }
            .abstract-title { font-weight: bold; }
            .two-column { column-count: 2; column-gap: 20px; text-align: justify; }
            .section-title { font-weight: bold; font-size: 11pt; text-align: center; margin-top: 15px; margin-bottom: 8px; }
            .content { text-align: justify; margin-bottom: 8px; white-space: pre-line; }
            .subsection { font-weight: bold; font-style: italic; margin-top: 10px; margin-bottom: 5px; }
            ul { margin: 8px 0; padding-left: 15px; }
            li { margin: 3px 0; }
            .algorithm { background: #f8f8f8; border: 1px solid #ddd; padding: 8px; margin: 8px 0; font-family: 'Courier New', monospace; font-size: 9pt; white-space: pre; }
            .references { font-size: 9pt; }
            .ref-item { margin: 3px 0; }
        </style>
        """
        
        # Build HTML
        html = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>{paper.title}</title>
    {css}
</head>
<body>
    <h1>{paper.title}</h1>
    
    <div class="authors">
"""
        
        # Authors
        for author in paper.authors:
            html += f"""
        <div class="author">
            <div class="author-name">{author.name}</div>
            <div class="author-affil">{author.affiliation}</div>
            <div class="author-email">{author.email}</div>
        </div>
"""
        
        html += """    </div>
    
    <div class="abstract">
        <span class="abstract-title">Abstract—</span>""" + paper.abstract + """
    </div>
"""
        
        if paper.doi:
            html += f'    <p><strong>DOI:</strong> {paper.doi}</p>\n'
        
        # Start two-column layout for body content
        html += '\n    <div class="two-column">\n'
        
        # Sections
        section_titles = {
            'introduction': 'I. INTRODUCTION',
            'literature_review': 'II. LITERATURE REVIEW',
            'methodology': 'III. METHODOLOGY',
            'results': 'IV. RESULTS',
            'discussion': 'V. DISCUSSION',
            'conclusion': 'VI. CONCLUSION',
            'references': 'REFERENCES'
        }
        
        for key, title in section_titles.items():
            if key in paper.sections:
                html += f'\n    <div class="section-title">{title}</div>\n'
                
                content = self._fix_encoding(paper.sections[key])
                # Convert bullets to HTML list items
                lines = content.split('\n')
                in_list = False
                
                for line in lines:
                    line = line.strip()
                    if not line:
                        if in_list:
                            html += '    </ul>\n'
                            in_list = False
                        continue
                    
                    # Check for bullet points
                    if line.startswith('•') or line.startswith('-') or line.startswith('*'):
                        if not in_list:
                            html += '    <ul>\n'
                            in_list = True
                        # Remove bullet character
                        item = line.lstrip('•-* ')
                        html += f'        <li>{item}</li>\n'
                    # Check for subsection headers (A. B. C. etc)
                    elif len(line) > 2 and line[0].isupper() and line[1] == '.' and line[2] == ' ':
                        if in_list:
                            html += '    </ul>\n'
                            in_list = False
                        html += f'    <div class="subsection">{line}</div>\n'
                    # Check for algorithm blocks
                    elif line.upper().startswith('ALGORITHM'):
                        if in_list:
                            html += '    </ul>\n'
                            in_list = False
                        html += f'    <div class="algorithm">{line}</div>\n'
                    else:
                        if in_list:
                            html += '    </ul>\n'
                            in_list = False
                        html += f'    <p class="content">{line}</p>\n'
                
                if in_list:
                    html += '    </ul>\n'
        
        # Close two-column div
        html += '    </div>\n'
        
        html += """
</body>
</html>"""
        
        return html

    def _create_styles(self):
        """Create PDF styles matching IEEE format"""
        styles = getSampleStyleSheet()
        
        # Title
        styles.add(ParagraphStyle(
            name='pdf_title',
            parent=styles['Heading1'],
            fontSize=24,
            leading=28,
            alignment=TA_CENTER,
            fontName='Times-Bold',
            spaceAfter=12
        ))
        
        # Author styles
        styles.add(ParagraphStyle(
            name='pdf_author_name',
            parent=styles['Normal'],
            fontSize=11,
            leading=13,
            alignment=TA_CENTER,
            fontName='Times-Bold'
        ))
        
        styles.add(ParagraphStyle(
            name='pdf_author_affil',
            parent=styles['Normal'],
            fontSize=10,
            leading=12,
            alignment=TA_CENTER,
            fontName='Times-Italic'
        ))
        
        styles.add(ParagraphStyle(
            name='pdf_author_email',
            parent=styles['Normal'],
            fontSize=9,
            leading=11,
            alignment=TA_CENTER,
            fontName='Times-Roman'
        ))
        
        # Section Heading (Centered)
        styles.add(ParagraphStyle(
            name='pdf_section',
            parent=styles['Heading2'],
            fontSize=10,
            leading=12,
            alignment=TA_CENTER,
            fontName='Times-Bold',
            spaceAfter=6,
            spaceBefore=12,
            textTransform='uppercase' 
        ))
        
        # Body Text
        styles.add(ParagraphStyle(
            name='pdf_body',
            parent=styles['BodyText'],
            fontSize=10,
            leading=12,
            alignment=TA_JUSTIFY,
            fontName='Times-Roman',
            firstLineIndent=12,
            spaceAfter=6,
            allowWidows=0,
            allowOrphans=0
        ))
        
        # Abstract
        styles.add(ParagraphStyle(
            name='pdf_abstract',
            parent=styles['BodyText'],
            fontSize=9,
            leading=11,
            alignment=TA_JUSTIFY,
            fontName='Times-Bold',
            firstLineIndent=0,
            spaceAfter=6,
            leftIndent=0,
            rightIndent=0
        ))
        
        # Caption
        styles.add(ParagraphStyle(
            name='pdf_caption',
            parent=styles['BodyText'],
            fontSize=8,
            alignment=TA_CENTER,
            fontName='Times-Italic',
            firstLineIndent=0
        ))
        
        return styles
    
    def _add_figures_to_pdf(self, figures, frame_width, styles):
        """Add figures and tables to PDF"""
        elements = []
        elements.append(Spacer(1, 0.1*inch))
        
        # Add figure images
        for key in sorted([k for k in figures.keys() if k.startswith('figure')]):
            try:
                fig_data = base64.b64decode(figures[key].data)
                fig_buf = BytesIO(fig_data)
                img = RLImage(fig_buf, width=frame_width*0.9, height=2.5*inch)
                elements.append(img)
                elements.append(Paragraph(figures[key].caption, styles['pdf_caption']))
                elements.append(Spacer(1, 0.1*inch))
            except Exception as e:
                print(f"[EXPORT] Figure error: {e}")
        
        # Add tables
        for key in sorted([k for k in figures.keys() if k.startswith('table')]):
            try:
                table_data = figures[key].data
                num_cols = len(table_data[0])
                col_width = frame_width / num_cols * 0.95
                
                # Wrap content in Paragraphs for text wrapping
                formatted_data = []
                for row in table_data:
                    formatted_row = []
                    for cell in row:
                        # Use a smaller font style for table content
                        p = Paragraph(str(cell), styles['pdf_caption']) 
                        formatted_row.append(p)
                    formatted_data.append(formatted_row)
                
                t = Table(formatted_data, colWidths=[col_width] * num_cols)
                t.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#667eea')),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                    ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                    ('LEFTPADDING', (0, 0), (-1, -1), 3),
                    ('RIGHTPADDING', (0, 0), (-1, -1), 3),
                    ('TOPPADDING', (0, 0), (-1, -1), 3),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 3),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ]))
                
                elements.append(t)
                elements.append(Paragraph(figures[key].caption, styles['pdf_caption']))
                elements.append(Spacer(1, 0.1*inch))
            except Exception as e:
                print(f"[EXPORT] Table error: {e}")
        
        return elements
